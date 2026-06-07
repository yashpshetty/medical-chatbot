import os
import sqlite3
import random
import time
import base64
import json
import hashlib
import requests
import wikipedia
import streamlit as st
from datetime import datetime, timedelta
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from PIL import Image
import io
import PyPDF2
import docx as python_docx

# ============================================================
# PAGE CONFIG
# ============================================================
st.set_page_config(
    page_title="Medical AI Assistant",
    page_icon="🩺",
    layout="centered"
)

# ============================================================
# CONFIGURATION
# ============================================================
HF_API_TOKEN  = os.environ.get("HF_API_TOKEN", "")
HF_API_URL    = "https://router.huggingface.co/v1/chat/completions"
MODEL         = "meta-llama/Llama-3.1-8B-Instruct"          # text model
VISION_MODEL  = "meta-llama/Llama-3.2-11B-Vision-Instruct"  # vision model for real image analysis
HF_HEADERS    = {"Authorization": f"Bearer {HF_API_TOKEN}"}

# Email config (set via env vars or Streamlit secrets)
SMTP_EMAIL    = os.environ.get("SMTP_EMAIL", "")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "")
SMTP_SERVER   = os.environ.get("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT     = int(os.environ.get("SMTP_PORT", "587"))

# ============================================================
# DATABASE SETUP
# ============================================================
DB_PATH = "medical_assistant.db"

def get_db():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            email       TEXT    UNIQUE NOT NULL,
            password    TEXT    NOT NULL,
            created_at  TEXT    DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS otp_store (
            email       TEXT PRIMARY KEY,
            otp         TEXT NOT NULL,
            expires_at  REAL NOT NULL
        );

        CREATE TABLE IF NOT EXISTS sessions (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            title       TEXT,
            created_at  TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        );

        CREATE TABLE IF NOT EXISTS messages (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id  INTEGER NOT NULL,
            role        TEXT NOT NULL,
            content     TEXT NOT NULL,
            image_data  TEXT,
            created_at  TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (session_id) REFERENCES sessions(id)
        );
    """)
    conn.commit()
    conn.close()

init_db()

# ============================================================
# PASSWORD HELPERS
# ============================================================
def hash_password(password: str) -> str:
    """SHA-256 hash (bcrypt-free fallback for portability)."""
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed

# ============================================================
# USER AUTH FUNCTIONS
# ============================================================
def user_exists(email: str) -> bool:
    conn = get_db()
    row = conn.execute("SELECT id FROM users WHERE email=?", (email,)).fetchone()
    conn.close()
    return row is not None

def register_user(email: str, password: str) -> bool:
    try:
        conn = get_db()
        conn.execute(
            "INSERT INTO users (email, password) VALUES (?,?)",
            (email, hash_password(password))
        )
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        return False

def login_user(email: str, password: str):
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
    conn.close()
    if row and verify_password(password, row["password"]):
        return dict(row)
    return None

def update_password(email: str, new_password: str):
    conn = get_db()
    conn.execute(
        "UPDATE users SET password=? WHERE email=?",
        (hash_password(new_password), email)
    )
    conn.commit()
    conn.close()

# ============================================================
# OTP FUNCTIONS
# ============================================================
def generate_otp() -> str:
    return str(random.randint(100000, 999999))

def store_otp(email: str, otp: str):
    conn = get_db()
    expires = time.time() + 600  # 10 minutes
    conn.execute(
        "INSERT OR REPLACE INTO otp_store (email, otp, expires_at) VALUES (?,?,?)",
        (email, otp, expires)
    )
    conn.commit()
    conn.close()

def verify_otp(email: str, otp: str) -> bool:
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM otp_store WHERE email=?", (email,)
    ).fetchone()
    conn.close()
    if row and row["otp"] == otp and time.time() < row["expires_at"]:
        return True
    return False

def delete_otp(email: str):
    conn = get_db()
    conn.execute("DELETE FROM otp_store WHERE email=?", (email,))
    conn.commit()
    conn.close()

def send_otp_email(email: str, otp: str) -> bool:
    """Send OTP via SMTP. Falls back to showing OTP on screen for demo."""
    if SMTP_EMAIL and SMTP_PASSWORD:
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart

            msg = MIMEMultipart("alternative")
            msg["Subject"] = "Your Medical AI Assistant OTP"
            msg["From"]    = SMTP_EMAIL
            msg["To"]      = email

            html = f"""
            <html><body>
            <h2>🩺 Medical AI Assistant</h2>
            <p>Your OTP for password reset is:</p>
            <h1 style="color:#2563eb;letter-spacing:6px">{otp}</h1>
            <p>This OTP is valid for <b>10 minutes</b>.</p>
            </body></html>
            """
            msg.attach(MIMEText(html, "html"))

            with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
                server.starttls()
                server.login(SMTP_EMAIL, SMTP_PASSWORD)
                server.sendmail(SMTP_EMAIL, email, msg.as_string())
            return True
        except Exception as e:
            st.warning(f"Email sending failed: {e}. Showing OTP here for demo.")
    # Demo mode: display OTP in the UI
    st.info(f"📧 **Demo Mode** — OTP for `{email}`: **{otp}**  *(In production, this is sent via email)*")
    return True

# ============================================================
# SESSION / CHAT DB FUNCTIONS
# ============================================================
def create_chat_session(user_id: int, title: str = "New Chat") -> int:
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO sessions (user_id, title) VALUES (?,?)",
        (user_id, title)
    )
    session_id = cur.lastrowid
    conn.commit()
    conn.close()
    return session_id

def get_user_sessions(user_id: int):
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM sessions WHERE user_id=? ORDER BY created_at DESC",
        (user_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def update_session_title(session_id: int, title: str):
    conn = get_db()
    conn.execute("UPDATE sessions SET title=? WHERE id=?", (title, session_id))
    conn.commit()
    conn.close()

def delete_session(session_id: int):
    conn = get_db()
    conn.execute("DELETE FROM messages WHERE session_id=?", (session_id,))
    conn.execute("DELETE FROM sessions WHERE id=?", (session_id,))
    conn.commit()
    conn.close()

def save_message(session_id: int, role: str, content: str, image_data: str = None):
    conn = get_db()
    conn.execute(
        "INSERT INTO messages (session_id, role, content, image_data) VALUES (?,?,?,?)",
        (session_id, role, content, image_data)
    )
    conn.commit()
    conn.close()

def get_session_messages(session_id: int):
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM messages WHERE session_id=? ORDER BY created_at ASC",
        (session_id,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# ============================================================
# WIKIPEDIA KNOWLEDGE BASE — cached once
# ============================================================
FALLBACK = {
    "Fever": "Fever is a temporary increase in body temperature above 38C, usually caused by infections. Treatment includes rest, fluids, and paracetamol or ibuprofen.",
    "Diabetes mellitus": "Diabetes is a condition of high blood glucose. Type 1 lacks insulin; Type 2 resists it. Managed with diet, exercise, and medications like metformin.",
    "Common cold": "The common cold is a viral upper respiratory infection. It resolves in 7–10 days with rest and fluids.",
    "Allergic disease": "Allergies are immune reactions to substances like pollen or food. Treated with antihistamines.",
}

TOPICS = [
    "Fever", "Cough", "Headache", "Diabetes mellitus",
    "Hypertension", "Common cold", "Influenza", "Sore throat",
    "Nausea", "Diarrhea", "Asthma", "Allergic disease",
    "Back pain", "Dehydration", "Anxiety"
]

@st.cache_resource(show_spinner="📚 Loading medical knowledge base...")
def load_knowledge_base():
    documents, loaded, failed = [], [], []
    for topic in TOPICS:
        try:
            text = wikipedia.summary(topic, sentences=6, auto_suggest=False)
            documents.append(text)
            loaded.append(topic)
        except Exception:
            if topic in FALLBACK:
                documents.append(FALLBACK[topic])
                loaded.append(f"{topic} (fallback)")
            else:
                failed.append(topic)
    vectorizer   = TfidfVectorizer()
    doc_vectors  = vectorizer.fit_transform(documents)
    return vectorizer, doc_vectors, documents, loaded, failed

vectorizer, doc_vectors, wiki_documents, loaded_topics, failed_topics = load_knowledge_base()

# ============================================================
# LLM / RAG CORE
# ============================================================
def query_llm(messages, max_tokens=500, model=None):
    if not HF_API_TOKEN:
        return "⚠️ HF_API_TOKEN not set. Add it to your environment variables or Streamlit secrets."
    payload = {"model": model or MODEL, "messages": messages, "max_tokens": max_tokens}
    try:
        resp   = requests.post(HF_API_URL, headers=HF_HEADERS, json=payload, timeout=60)
        result = resp.json()
        if isinstance(result, dict) and "error" in result:
            return f"API Error: {result['error']}"
        return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"Error contacting LLM: {str(e)}"

def retrieve_context(query: str) -> str:
    q_vec  = vectorizer.transform([query])
    scores = cosine_similarity(q_vec, doc_vectors)
    return wiki_documents[scores.argmax()]

def image_to_base64(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()

def classify_intent(user_input: str, chat_history: list) -> str:
    history_text = ""
    if chat_history:
        history_text = "\n\nRecent conversation:\n"
        for turn in chat_history[-4:]:
            history_text += f"User: {turn['user']}\nAssistant: {turn['bot']}\n"
    messages = [
        {
            "role": "system",
            "content": (
                "You are an intent classifier. Classify the user message into ONE of:\n"
                "- greeting\n- medical\n- non_medical\n\n"
                "If the assistant previously asked a follow-up medical question, "
                "classify the reply as 'medical'.\n"
                f"{history_text}\n"
                "Reply with ONLY one word: greeting, medical, or non_medical."
            )
        },
        {"role": "user", "content": f"Classify: {user_input}"}
    ]
    result = query_llm(messages, max_tokens=10).strip().lower()
    if "greeting" in result:
        return "greeting"
    if "non_medical" in result or "non-medical" in result:
        return "non_medical"
    return "medical"

def rag_query(user_input: str, chat_history: list) -> str:
    context  = retrieve_context(user_input)
    messages = [
        {
            "role": "system",
            "content": (
                "You are a helpful Medical AI Assistant. Use the context to answer.\n"
                "Be concise, empathetic, and medically responsible.\n"
                "If something sounds serious, advise the user to consult a doctor.\n\n"
                f"Medical Context:\n{context}"
            )
        }
    ]
    for turn in chat_history:
        messages.append({"role": "user",      "content": turn["user"]})
        messages.append({"role": "assistant", "content": turn["bot"]})
    messages.append({"role": "user", "content": user_input})
    return query_llm(messages, max_tokens=500)

def analyze_medical_image(image: Image.Image, user_question: str, chat_history: list) -> str:
    """Analyse uploaded image using LLaMA 3.2 Vision — sends actual pixel data."""
    img_b64  = image_to_base64(image)
    context  = retrieve_context(user_question)

    system_msg = {
        "role": "system",
        "content": (
            "You are a Medical AI Assistant with vision capabilities. "
            "The user has uploaded a real medical image (skin condition, rash, wound, "
            "x-ray scan, lab report photo, prescription scan, or other visible symptom). "
            "Carefully examine the image and provide helpful, accurate medical information "
            "based on what you actually see. Highlight any visible abnormalities, "
            "patterns, or notable features. Always advise consulting a qualified doctor "
            "for proper diagnosis and treatment.\n\n"
            f"Relevant Medical Knowledge:\n{context}"
        )
    }

    # Build history context (text only — keep context window lean)
    messages = [system_msg]
    for turn in chat_history[-3:]:
        messages.append({"role": "user",      "content": turn["user"]})
        messages.append({"role": "assistant", "content": turn["bot"]})

    # Vision message: image + question as a multi-part content block
    messages.append({
        "role": "user",
        "content": [
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{img_b64}"}
            },
            {
                "type": "text",
                "text": user_question if user_question.strip()
                        else "Please examine this medical image and describe what you see in detail. "
                             "Highlight any visible abnormalities or areas of concern."
            }
        ]
    })

    return query_llm(messages, max_tokens=600, model=VISION_MODEL)

# ── DOCUMENT EXTRACTION ───────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file."""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages  = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i+1}]\n{text.strip()}")
        if not pages:
            return ""
        return "\n\n".join(pages)
    except Exception as e:
        return f"[PDF read error: {e}]"

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file."""
    try:
        doc   = python_docx.Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        return f"[DOCX read error: {e}]"

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode plain text / CSV files."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")

def extract_document_text(uploaded_file) -> tuple[str, str]:
    """
    Read an uploaded Streamlit file object.
    Returns (extracted_text, file_type_label).
    """
    name  = uploaded_file.name.lower()
    raw   = uploaded_file.read()
    uploaded_file.seek(0)          # reset so caller can re-read if needed

    if name.endswith(".pdf"):
        return extract_text_from_pdf(raw), "PDF"
    elif name.endswith(".docx"):
        return extract_text_from_docx(raw), "DOCX"
    elif name.endswith(".doc"):
        return "[Legacy .doc format not supported. Please save as .docx and re-upload.]", "DOC"
    else:
        # .txt / .csv / .rtf / anything text-like
        return extract_text_from_txt(raw), "TXT"

def analyze_medical_document(doc_text: str, file_name: str,
                              user_question: str, chat_history: list) -> str:
    """Analyse extracted medical document text and answer the user's question."""
    # Truncate to ~3000 chars to stay within LLM context
    MAX_CHARS = 3000
    truncated = False
    if len(doc_text) > MAX_CHARS:
        doc_text  = doc_text[:MAX_CHARS]
        truncated = True

    context = retrieve_context(user_question or doc_text[:200])

    system_prompt = (
        "You are a Medical AI Assistant specialised in reading and explaining medical documents.\n"
        "The user has uploaded a medical document (lab report, prescription, discharge summary, "
        "test results, radiology report, etc.).\n\n"
        "Your tasks:\n"
        "1. Summarise the key findings from the document in plain language.\n"
        "2. Highlight any abnormal values, concerning terms, or important instructions.\n"
        "3. Answer the user's specific question about the document.\n"
        "4. Always advise the user to consult their doctor for interpretation and treatment.\n\n"
        f"Relevant Medical Knowledge:\n{context}\n\n"
        f"--- DOCUMENT CONTENT ({file_name}) ---\n{doc_text}\n"
        + ("...[document truncated for length]" if truncated else "")
    )

    messages = [{"role": "system", "content": system_prompt}]
    for turn in chat_history[-3:]:
        messages.append({"role": "user",      "content": turn["user"]})
        messages.append({"role": "assistant", "content": turn["bot"]})

    question = user_question if user_question.strip() else "Please summarise this medical document and highlight anything important."
    messages.append({"role": "user", "content": question})

    return query_llm(messages, max_tokens=600)

def follow_up_with_context(user_input: str, prior_analysis: str, chat_history: list) -> str:
    """Answer a follow-up question using the cached first analysis as context."""
    messages = [
        {
            "role": "system",
            "content": (
                "You are a Medical AI Assistant. The user has already had their medical file analysed. "
                "Below is the initial analysis you provided. Use it as context to answer the user's "
                "follow-up question — do NOT re-analyse or re-summarise the file from scratch. "
                "Just answer the specific follow-up question concisely.\n\n"
                f"Initial analysis:\n{prior_analysis}"
            )
        }
    ]
    for turn in chat_history[-4:]:
        messages.append({"role": "user",      "content": turn["user"]})
        messages.append({"role": "assistant", "content": turn["bot"]})
    messages.append({"role": "user", "content": user_input})
    return query_llm(messages, max_tokens=500)


def medical_bot(user_input: str, chat_history: list,
                uploaded_image: Image.Image = None,
                doc_text: str = None, doc_name: str = None) -> str:
    """
    Route the request.
    - First upload of a doc/image  → full analysis, result cached in session_state
    - Follow-up with same doc/image → answer using cached analysis, no re-processing
    - Plain question                → RAG or greeting/non-medical routing
    """
    # ── DOCUMENT ────────────────────────────────────────────────
    if doc_text is not None:
        cache = st.session_state.get("doc_analysis_cache")
        # Use cache only if it's for the same document
        if cache and cache.get("doc_name") == doc_name:
            return follow_up_with_context(user_input, cache["summary"], chat_history)
        # First time — full analysis, then cache it
        result = analyze_medical_document(doc_text, doc_name or "document", user_input, chat_history)
        st.session_state["doc_analysis_cache"] = {"doc_name": doc_name, "summary": result}
        return result

    # ── IMAGE ────────────────────────────────────────────────────
    if uploaded_image is not None:
        cache = st.session_state.get("image_analysis_cache")
        img_name = st.session_state.get("staged_image_name", "")
        if cache and cache.get("image_name") == img_name:
            return follow_up_with_context(user_input, cache["summary"], chat_history)
        # First time — full vision analysis, then cache it
        result = analyze_medical_image(uploaded_image, user_input, chat_history)
        st.session_state["image_analysis_cache"] = {"image_name": img_name, "summary": result}
        return result

    # ── PLAIN TEXT ───────────────────────────────────────────────
    intent = classify_intent(user_input, chat_history)
    if intent == "greeting":
        return ("Hello! I am your Medical AI Assistant 🩺  "
                "How can I help you today? Please describe your symptoms "
                "or ask a health-related question.")
    if intent == "non_medical":
        return "I am a Medical AI Assistant. Please ask health-related questions only."
    return rag_query(user_input, chat_history)

# ============================================================
# SESSION STATE DEFAULTS
# ============================================================
def defaults():
    keys = {
        "auth_page":          "login",       # login | register | forgot | verify_otp | reset
        "logged_in":          False,
        "user":               None,
        "active_session_id":  None,
        "messages":           [],            # display list
        "chat_history":       [],            # LLM memory list [{user, bot}]
        "fp_email":           "",            # forgot-password flow
        "fp_otp":             "",
        "pending_input":      None,
        # ── file caches ──────────────────────────────────────────
        "doc_analysis_cache":   None,
        "image_analysis_cache": None,
        "staged_image":         None,
        "staged_image_name":    "",
        "staged_doc_text":      None,
        "staged_doc_name":      "",
    }
    for k, v in keys.items():
        if k not in st.session_state:
            st.session_state[k] = v

defaults()

# ============================================================
# HELPER: load DB session into session_state
# ============================================================
def load_db_session(session_id: int):
    msgs = get_session_messages(session_id)
    display, history = [], []
    for m in msgs:
        if m["role"] == "user":
            display.append({"role": "user",      "content": m["content"],
                            "image_data": m.get("image_data")})
        else:
            display.append({"role": "assistant", "content": m["content"]})
    # rebuild LLM memory (user+bot pairs)
    user_msgs = [m for m in msgs if m["role"] == "user"]
    bot_msgs  = [m for m in msgs if m["role"] == "assistant"]
    for u, b in zip(user_msgs, bot_msgs):
        history.append({"user": u["content"], "bot": b["content"]})
    st.session_state.messages      = display
    st.session_state.chat_history  = history
    st.session_state.active_session_id = session_id

# ============================================================
# CUSTOM CSS
# ============================================================
st.markdown("""
<style>
/* ── Auth card ── */
.auth-card {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 16px;
    padding: 2rem 2.2rem;
    max-width: 420px;
    margin: 3rem auto;
    box-shadow: 0 4px 24px rgba(0,0,0,.07);
}
.auth-title {
    text-align: center;
    font-size: 1.6rem;
    font-weight: 700;
    color: #1e40af;
    margin-bottom: .2rem;
}
.auth-sub {
    text-align: center;
    color: #64748b;
    margin-bottom: 1.4rem;
    font-size: .92rem;
}
/* ── Chat bubbles ── */
[data-testid="stChatMessage"] {
    border-radius: 12px !important;
    margin-bottom: .4rem !important;
}
/* ── Sidebar sessions ── */
.session-card {
    background: #f1f5f9;
    border-radius: 10px;
    padding: .55rem .75rem;
    margin-bottom: .45rem;
    cursor: pointer;
    border-left: 3px solid transparent;
    transition: border-color .15s;
}
.session-card:hover { border-left-color: #3b82f6; }
.session-active   { border-left-color: #1d4ed8 !important; background: #dbeafe; }
</style>
""", unsafe_allow_html=True)

# ============================================================
# AUTH PAGES
# ============================================================

def page_login():
    st.markdown("""
    <div class='auth-card'>
      <div class='auth-title'>🩺 Medical AI Assistant</div>
      <div class='auth-sub'>Sign in to your account</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form("login_form"):
        email    = st.text_input("📧 Email address", placeholder="you@example.com")
        password = st.text_input("🔒 Password",       type="password")
        submit   = st.form_submit_button("Sign In", use_container_width=True, type="primary")

    if submit:
        if not email or not password:
            st.error("Please fill in all fields.")
        else:
            user = login_user(email.strip().lower(), password)
            if user:
                st.session_state.logged_in = True
                st.session_state.user      = user
                st.success(f"Welcome back, {user['email']} 👋")
                time.sleep(0.6)
                st.rerun()
            else:
                st.error("Invalid email or password.")

    col1, col2 = st.columns(2)
    if col1.button("🆕 Create Account", use_container_width=True):
        st.session_state.auth_page = "register"
        st.rerun()
    if col2.button("🔑 Forgot Password", use_container_width=True):
        st.session_state.auth_page = "forgot"
        st.rerun()


def page_register():
    st.markdown("""
    <div class='auth-card'>
      <div class='auth-title'>🩺 Create Account</div>
      <div class='auth-sub'>Register to get started</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form("reg_form"):
        email    = st.text_input("📧 Email address", placeholder="you@example.com")
        password = st.text_input("🔒 Password (min 6 chars)", type="password")
        confirm  = st.text_input("🔒 Confirm Password",       type="password")
        submit   = st.form_submit_button("Register", use_container_width=True, type="primary")

    if submit:
        email = email.strip().lower()
        if not email or not password or not confirm:
            st.error("Please fill in all fields.")
        elif len(password) < 6:
            st.error("Password must be at least 6 characters.")
        elif password != confirm:
            st.error("Passwords do not match.")
        elif user_exists(email):
            st.error("An account with this email already exists.")
        else:
            if register_user(email, password):
                st.success("Account created! Please sign in.")
                time.sleep(1)
                st.session_state.auth_page = "login"
                st.rerun()
            else:
                st.error("Registration failed. Try again.")

    if st.button("← Back to Sign In", use_container_width=True):
        st.session_state.auth_page = "login"
        st.rerun()


def page_forgot_password():
    st.markdown("""
    <div class='auth-card'>
      <div class='auth-title'>🔑 Forgot Password</div>
      <div class='auth-sub'>Enter your email to receive a 6-digit OTP</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form("fp_form"):
        email  = st.text_input("📧 Registered Email", placeholder="you@example.com")
        submit = st.form_submit_button("Send OTP", use_container_width=True, type="primary")

    if submit:
        email = email.strip().lower()
        if not email:
            st.error("Please enter your email address.")
        elif not user_exists(email):
            st.error("No account found with this email.")
        else:
            otp = generate_otp()
            store_otp(email, otp)
            send_otp_email(email, otp)
            st.session_state.fp_email  = email
            st.session_state.auth_page = "verify_otp"
            st.rerun()

    if st.button("← Back to Sign In", use_container_width=True):
        st.session_state.auth_page = "login"
        st.rerun()


def page_verify_otp():
    st.markdown("""
    <div class='auth-card'>
      <div class='auth-title'>✉️ Verify OTP</div>
      <div class='auth-sub'>Enter the 6-digit code sent to your email</div>
    </div>
    """, unsafe_allow_html=True)

    email = st.session_state.get("fp_email", "")
    st.info(f"OTP sent to: **{email}**  (valid for 10 minutes)")

    with st.form("otp_form"):
        otp    = st.text_input("🔢 Enter OTP", max_chars=6, placeholder="123456")
        submit = st.form_submit_button("Verify OTP", use_container_width=True, type="primary")

    if submit:
        if verify_otp(email, otp.strip()):
            st.session_state.fp_otp    = otp.strip()
            st.session_state.auth_page = "reset"
            st.rerun()
        else:
            st.error("Invalid or expired OTP. Please try again.")

    col1, col2 = st.columns(2)
    if col1.button("🔄 Resend OTP", use_container_width=True):
        new_otp = generate_otp()
        store_otp(email, new_otp)
        send_otp_email(email, new_otp)
        st.success("New OTP sent!")
    if col2.button("← Back", use_container_width=True):
        st.session_state.auth_page = "forgot"
        st.rerun()


def page_reset_password():
    st.markdown("""
    <div class='auth-card'>
      <div class='auth-title'>🔒 Reset Password</div>
      <div class='auth-sub'>Enter your new password</div>
    </div>
    """, unsafe_allow_html=True)

    email = st.session_state.get("fp_email", "")
    with st.form("reset_form"):
        new_pw  = st.text_input("🔒 New Password (min 6 chars)", type="password")
        confirm = st.text_input("🔒 Confirm New Password",       type="password")
        submit  = st.form_submit_button("Reset Password", use_container_width=True, type="primary")

    if submit:
        if not new_pw or not confirm:
            st.error("Please fill in both fields.")
        elif len(new_pw) < 6:
            st.error("Password must be at least 6 characters.")
        elif new_pw != confirm:
            st.error("Passwords do not match.")
        else:
            update_password(email, new_pw)
            delete_otp(email)
            st.success("Password reset successfully! Please sign in.")
            time.sleep(1.2)
            st.session_state.auth_page = "login"
            st.session_state.fp_email  = ""
            st.session_state.fp_otp    = ""
            st.rerun()

# ============================================================
# MAIN CHAT UI (authenticated)
# ============================================================

def page_chat():
    user    = st.session_state.user
    user_id = user["id"]

    # ── SIDEBAR ──────────────────────────────────────────────
    with st.sidebar:
        st.markdown(f"### 👤 {user['email']}")
        if st.button("🚪 Sign Out", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            defaults()
            st.rerun()

        st.markdown("---")

        # New Chat button
        if st.button("➕ New Chat", use_container_width=True, type="primary"):
            st.session_state.active_session_id    = None
            st.session_state.messages             = []
            st.session_state.chat_history         = []
            st.session_state.doc_analysis_cache   = None
            st.session_state.image_analysis_cache = None
            st.session_state.pop("staged_image",    None)
            st.session_state.pop("staged_image_name", None)
            st.session_state.pop("staged_doc_text", None)
            st.session_state.pop("staged_doc_name", None)
            st.rerun()

        st.markdown("### 🕒 Chat History")

        sessions = get_user_sessions(user_id)
        if not sessions:
            st.info("No previous chats.\nStart a new conversation!")
        else:
            for s in sessions:
                active = (st.session_state.active_session_id == s["id"])
                card_class = "session-card session-active" if active else "session-card"
                label = s["title"] or "Untitled"
                label_short = label[:32] + "…" if len(label) > 32 else label
                date_str = s["created_at"][:10]

                with st.container():
                    col_title, col_del = st.columns([5, 1])
                    with col_title:
                        st.markdown(
                            f"<div class='{card_class}'>"
                            f"<b>{label_short}</b><br>"
                            f"<small style='color:#64748b'>{date_str}</small>"
                            f"</div>",
                            unsafe_allow_html=True
                        )
                        if st.button(
                            f"↩ Open##{s['id']}",
                            key=f"open_{s['id']}",
                            use_container_width=True
                        ):
                            load_db_session(s["id"])
                            st.session_state.doc_analysis_cache   = None
                            st.session_state.image_analysis_cache = None
                            st.rerun()
                    with col_del:
                        if st.button("🗑", key=f"del_{s['id']}"):
                            delete_session(s["id"])
                            if st.session_state.active_session_id == s["id"]:
                                st.session_state.active_session_id    = None
                                st.session_state.messages             = []
                                st.session_state.chat_history         = []
                                st.session_state.doc_analysis_cache   = None
                                st.session_state.image_analysis_cache = None
                            st.rerun()

    # ── MAIN AREA ─────────────────────────────────────────────
    st.title("🩺 Medical AI Assistant")
    st.caption("Powered by LLaMA 3.1 + RAG Knowledge Base")
    st.warning(
        "This chatbot is for **informational purposes only** and is **not** "
        "a substitute for professional medical advice."
    )

    # ── DISPLAY MESSAGES ──────────────────────────────────────
    for msg in st.session_state.messages:
        avatar = "🧑" if msg["role"] == "user" else "🩺"
        with st.chat_message(msg["role"], avatar=avatar):
            if msg.get("image_data"):
                try:
                    img_bytes = base64.b64decode(msg["image_data"])
                    st.image(Image.open(io.BytesIO(img_bytes)), caption="Attached image", width=200)
                except Exception:
                    pass
            if msg.get("doc_name"):
                st.markdown(
                    f"<span style='background:#dbeafe;color:#1e40af;"
                    f"padding:3px 10px;border-radius:20px;font-size:0.82rem;'>"
                    f"📄 {msg['doc_name']}</span>",
                    unsafe_allow_html=True
                )
            st.markdown(msg["content"])

    # ── EXAMPLE BUTTONS (empty chat) ─────────────────────────
    if not st.session_state.messages:
        st.markdown("#### 💡 Try asking:")
        examples = [
            "I have a fever and headache since yesterday",
            "What are the symptoms of diabetes?",
            "I have a dry cough and sore throat",
            "How to manage high blood pressure?",
        ]
        cols = st.columns(2)
        for i, ex in enumerate(examples):
            if cols[i % 2].button(ex, use_container_width=True, key=f"ex_{i}"):
                st.session_state.pending_input = ex
                st.rerun()

    # ── SEND MESSAGE HELPER ───────────────────────────────────
    def send_message(user_text: str,
                     image: Image.Image = None,
                     doc_text: str = None, doc_name: str = None):
        if not user_text.strip() and doc_text:
            user_text = "Please summarise this medical document and highlight anything important."

        if not st.session_state.active_session_id:
            title = user_text[:40] + ("…" if len(user_text) > 40 else "")
            sid   = create_chat_session(user_id, title)
            st.session_state.active_session_id = sid
        else:
            sid = st.session_state.active_session_id

        img_b64 = image_to_base64(image) if image is not None else None

        st.session_state.messages.append({
            "role": "user", "content": user_text,
            "image_data": img_b64, "doc_name": doc_name
        })
        extra = json.dumps({"doc_name": doc_name}) if doc_name else None
        save_message(sid, "user", user_text, img_b64 or extra)

        with st.chat_message("user", avatar="🧑"):
            if img_b64:
                st.image(Image.open(io.BytesIO(base64.b64decode(img_b64))),
                         caption="Attached image", width=200)
            if doc_name:
                st.markdown(
                    f"<span style='background:#dbeafe;color:#1e40af;"
                    f"padding:3px 10px;border-radius:20px;font-size:0.82rem;'>"
                    f"📄 {doc_name}</span>", unsafe_allow_html=True)
            st.markdown(user_text)

        with st.chat_message("assistant", avatar="🩺"):
            with st.spinner("Analysing..." if (doc_text or image) else "Thinking..."):
                response = medical_bot(
                    user_text, st.session_state.chat_history,
                    uploaded_image=image, doc_text=doc_text, doc_name=doc_name
                )
            st.markdown(response)

        st.session_state.messages.append({"role": "assistant", "content": response})
        st.session_state.chat_history.append({"user": user_text, "bot": response})
        save_message(sid, "assistant", response)

        sessions = get_user_sessions(user_id)
        this_s   = next((s for s in sessions if s["id"] == sid), None)
        if this_s and this_s["title"] in ("New Chat", ""):
            update_session_title(sid, user_text[:40])

        # Clear staged file after sending (keep analysis caches for follow-ups)
        st.session_state.pop("staged_image",      None)
        st.session_state.pop("staged_image_name", None)
        st.session_state.pop("staged_doc_text",   None)
        st.session_state.pop("staged_doc_name",   None)

    # ── HANDLE EXAMPLE BUTTON CLICK ──────────────────────────
    if st.session_state.pending_input:
        pi = st.session_state.pending_input
        st.session_state.pending_input = None
        send_message(pi)
        st.rerun()

    # ── UNIFIED FILE UPLOAD + CHAT INPUT AT BOTTOM ───────────
    staged_image    = st.session_state.get("staged_image",    None)
    staged_doc_text = st.session_state.get("staged_doc_text", None)
    staged_doc_name = st.session_state.get("staged_doc_name", None)

    # Show preview of whatever is staged
    if staged_image:
        col_img, col_x = st.columns([9, 1])
        with col_img:
            st.image(staged_image,
                     caption=f"📷 {st.session_state.get('staged_image_name','')}",
                     width=150)
        with col_x:
            st.write("")
            if st.button("✕", key="rm_img", help="Remove image"):
                st.session_state.pop("staged_image",      None)
                st.session_state.pop("staged_image_name", None)
                st.rerun()

    if staged_doc_name:
        col_doc, col_xd = st.columns([9, 1])
        with col_doc:
            st.markdown(
                f"<div style='background:#f1f5f9;border-radius:8px;padding:6px 12px;"
                f"font-size:0.85rem;color:#1e40af;'>📄 <b>{staged_doc_name}</b> ready</div>",
                unsafe_allow_html=True)
        with col_xd:
            st.write("")
            if st.button("✕", key="rm_doc", help="Remove document"):
                st.session_state.pop("staged_doc_text", None)
                st.session_state.pop("staged_doc_name", None)
                st.rerun()

    # Single uploader — accepts images AND documents
    uploaded_file = st.file_uploader(
        "📎 Attach image or document",
        type=["jpg", "jpeg", "png", "webp", "bmp", "pdf", "docx", "txt", "csv"],
        key="unified_upload",
        label_visibility="visible",
        help="Upload an image or document, then type your question and press Enter"
    )

    if uploaded_file:
        ext = uploaded_file.name.lower().rsplit(".", 1)[-1]
        if ext in ("jpg", "jpeg", "png", "webp", "bmp"):
            pil = Image.open(uploaded_file)
            st.session_state["staged_image"]      = pil
            st.session_state["staged_image_name"] = uploaded_file.name
            st.session_state.pop("staged_doc_text", None)
            st.session_state.pop("staged_doc_name", None)
        else:
            with st.spinner(f"Reading {uploaded_file.name}…"):
                doc_text, doc_type = extract_document_text(uploaded_file)
            if doc_text.strip() and not (doc_text.startswith("[") and "error" in doc_text.lower()):
                st.session_state["staged_doc_text"] = doc_text
                st.session_state["staged_doc_name"] = uploaded_file.name
                st.session_state.pop("staged_image",      None)
                st.session_state.pop("staged_image_name", None)
            else:
                st.error(f"Could not read file: {doc_text}")
        st.rerun()

    # Chat input placeholder changes based on what's staged
    if staged_doc_text:
        placeholder = "Ask about the document, or press Enter for a summary…"
    elif staged_image:
        placeholder = "Ask a question about the image, then press Enter…"
    else:
        placeholder = "Describe your symptoms or ask a health question…"

    user_input = st.chat_input(placeholder)

    if user_input is not None:
        active_image    = st.session_state.get("staged_image")
        active_doc_text = st.session_state.get("staged_doc_text")
        active_doc_name = st.session_state.get("staged_doc_name")
        if user_input.strip() or active_doc_text or active_image:
            send_message(
                user_input,
                image=active_image,
                doc_text=active_doc_text,
                doc_name=active_doc_name
            )
            st.rerun()


# ============================================================
# ROUTER
# ============================================================
if not st.session_state.logged_in:
    page = st.session_state.auth_page
    if page == "register":
        page_register()
    elif page == "forgot":
        page_forgot_password()
    elif page == "verify_otp":
        page_verify_otp()
    elif page == "reset":
        page_reset_password()
    else:
        page_login()
else:
    page_chat()
